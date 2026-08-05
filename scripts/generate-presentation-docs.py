#!/usr/bin/env python3
"""Generate presentation DOCX files for Uziel and Erick."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "presentation"


def read_snippet(rel_path: str, start: int, end: int) -> str:
    lines = (ROOT / rel_path).read_text(encoding="utf-8").splitlines()
    return "\n".join(lines[start - 1 : end])


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    if "CodeBlock" not in [s.name for s in doc.styles]:
        style = doc.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Consolas"
        style.font.size = Pt(9)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    doc.add_heading(title, 0)
    p = doc.add_paragraph(subtitle)
    p.runs[0].italic = True
    doc.add_paragraph()


def add_h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_code(doc: Document, caption: str, code: str) -> None:
    add_p(doc, caption)
    for line in code.splitlines():
        p = doc.add_paragraph(line)
        if "CodeBlock" in [s.name for s in doc.styles]:
            p.style = doc.styles["CodeBlock"]
        for run in p.runs:
            run.font.name = "Consolas"
            run.font.size = Pt(8)
    doc.add_paragraph()


def add_page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build_uziel_doc() -> Document:
    doc = Document()
    setup_styles(doc)
    add_title(
        doc,
        "manage-system — Half A: Application Core",
        "Presenter: Uziel Almonte | Weeks 1–4 | Build the product users interact with",
    )

    add_h(doc, "1. Executive Summary")
    add_p(
        doc,
        "This document covers the application core of the inventory management system: Docker infrastructure, "
        "PostgreSQL with Alembic migrations, the Flask REST API (products, stock, audit, reports), Keycloak-based "
        "authentication and RBAC, and the Jinja2 + HTMX frontend. Erick's half (Weeks 5–6) validates, observes, "
        "and delivers this application through automated testing, observability, and CI/CD.",
    )

    add_h(doc, "2. TODO Days Covered (Uziel)")
    add_bullets(
        doc,
        [
            "Week 1 (Days 1–7): Repo, Docker Compose, Flask structure, Alembic migrations",
            "Week 2 (Days 8–12): Stock movements, audit API, OpenAPI/Swagger, reports endpoints",
            "Week 3 (Days 15–21): Keycloak realm, JWT middleware, @require_scope, OAuth2 login/logout",
            "Week 4 (Days 23–28): Dashboard, products/stock/audit/reports UI, scope-gated navigation",
            "Co-credit: Day 22 base layout, Day 13 unit tests, audit listeners (Erick De la Rosa)",
        ],
    )

    add_h(doc, "3. Infrastructure and Environment Variables")
    add_p(
        doc,
        "Why: We containerize the stack so every developer and CI runner gets identical Postgres, Keycloak, and Flask "
        "environments. Secrets live in .env, never in code.",
    )
    add_p(doc, "Key variables and where they come from:")
    add_bullets(
        doc,
        [
            "POSTGRES_USER / POSTGRES_PASSWORD / POSTGRES_DB — defined in .env, consumed by docker-compose.yml db service",
            "DATABASE_URL — postgresql://user:pass@db:5432/inventory_db (host db is the Docker network name)",
            "FLASK_APP=app.main — tells Flask CLI which module creates the application",
            "KEYCLOAK_REALM_URL — http://localhost:8080/realms/inventory-realm (browser + JWKS fetch)",
            "KEYCLOAK_TOKEN_URL — http://keycloak:8080/... (internal Docker hostname for server-to-server)",
        ],
    )
    add_code(
        doc,
        "docker-compose.yml — core services (excerpt):",
        read_snippet("docker-compose.yml", 4, 20),
    )

    add_page_break(doc)
    add_h(doc, "4. Products API")
    add_p(
        doc,
        "Why: Products are the central entity. Every stock movement references a product. The API enforces SKU "
        "uniqueness, pagination, search, and scope-based access (product:view vs product:manage).",
    )
    add_code(
        doc,
        "app/products/views.py — list products with scope guard:",
        read_snippet("app/products/views.py", 35, 68),
    )
    add_code(
        doc,
        "app/products/views.py — create product:",
        read_snippet("app/products/views.py", 70, 98),
    )

    add_h(doc, "5. Stock API")
    add_p(
        doc,
        "Why: Inventory changes must be traceable. Every entry/exit creates a StockMovement row and updates product.qty "
        "atomically. Exit movements reject insufficient stock.",
    )
    add_code(
        doc,
        "app/stock/views.py — movement registration:",
        read_snippet("app/stock/views.py", 21, 67),
    )
    add_p(
        doc,
        "Variables: product_id, type (entry|exit), qty_change come from JSON body validated by StockMovementSchema. "
        "user defaults to 'system' for API calls; the UI pre-fills the logged-in username.",
    )

    add_h(doc, "5b. Audit Trail (Day 10)")
    add_p(
        doc,
        "Why: Regulatory and operational requirements demand knowing who changed what. SQLAlchemy event listeners "
        "automatically write to audit_logs on INSERT, UPDATE, DELETE for products and stock_movements.",
    )
    add_code(
        doc,
        "app/audit/listeners.py — capture INSERT/UPDATE:",
        read_snippet("app/audit/listeners.py", 26, 73),
    )
    add_p(
        doc,
        "The user field comes from Flask g.user, set by login_required from session['user'].preferred_username.",
    )

    add_h(doc, "5c. UI Routes and Scope Enforcement (app/main.py)")
    add_p(
        doc,
        "Every HTML route uses @login_required and @require_ui_scope. Direct URL access without permission "
        "redirects to dashboard with flash 'Acceso denegado'.",
    )
    add_bullets(
        doc,
        [
            "/ — dashboard (login_required)",
            "/products — product:view",
            "/products/new — product:manage",
            "/stock — stock:view",
            "/stock/update — stock:manage",
            "/audit — audit:view",
            "/reports — report:view",
        ],
    )

    add_page_break(doc)
    add_h(doc, "6. Keycloak and RBAC")
    add_p(
        doc,
        "Why: Role-Based Access Control separates employees (alice_worker) from managers (kratos_boss). Permissions "
        "are fine-grained scopes stored in Keycloak realm roles, not hardcoded in Flask.",
    )
    add_bullets(
        doc,
        [
            "alice_worker / password123 — group:employee — product:view, stock:view, stock:manage",
            "kratos_boss / password123 — group:manager — all scopes including audit:view, report:view",
        ],
    )
    add_p(doc, "OAuth2 browser flow:")
    add_bullets(
        doc,
        [
            "1. User visits /auth/login-page → /auth/login",
            "2. Flask redirects browser to Keycloak (KEYCLOAK_AUTHORIZE_URL)",
            "3. User authenticates at Keycloak",
            "4. Keycloak redirects to /auth/callback with authorization code",
            "5. Flask exchanges code for tokens (KEYCLOAK_TOKEN_URL, internal Docker host)",
            "6. Callback decodes access_token JWT, extracts realm_access.roles containing ':' as user_scopes",
            "7. session['user'], session['token'], session['user_scopes'] stored; redirect to /",
        ],
    )
    add_code(
        doc,
        "app/auth/views.py — callback stores scopes from JWT:",
        read_snippet("app/auth/views.py", 41, 74),
    )
    add_code(
        doc,
        "app/auth/middleware.py — API JWT validation:",
        read_snippet("app/auth/middleware.py", 33, 50),
    )
    add_code(
        doc,
        "app/auth/middleware.py — require_scope:",
        read_snippet("app/auth/middleware.py", 92, 103),
    )
    add_code(
        doc,
        "app/auth/middleware.py — require_ui_scope:",
        read_snippet("app/auth/middleware.py", 127, 145),
    )

    add_page_break(doc)
    add_h(doc, "7. Frontend (Jinja2 + HTMX)")
    add_p(
        doc,
        "Why: Server-rendered HTML with HTMX gives dynamic UX without a separate React/Vue build step. "
        "TailwindCSS via CDN provides responsive styling.",
    )
    add_code(
        doc,
        "app/templates/base.html — scope-gated nav:",
        read_snippet("app/templates/base.html", 31, 47),
    )

    add_h(doc, "8. Role-Based Walkthrough (Demo Script)")
    add_bullets(
        doc,
        [
            "docker compose up -d && docker compose exec web flask db upgrade",
            "Login as alice_worker — dashboard, products, stock; NO audit nav",
            "Login as kratos_boss — create product, stock entry/exit, audit, reports",
            "Open http://localhost:5000/docs — Swagger UI",
        ],
    )

    add_h(doc, "9. Bridge to Erick's Half")
    add_p(
        doc,
        "The application core exposes REST API at /api/*, UI at /*, /metrics, and openapi_spec.json. "
        "Erick's half runs tests, observability, and CI/CD against this foundation.",
    )

    add_h(doc, "Appendix: Commands")
    add_bullets(
        doc,
        [
            "Start stack: docker compose up -d --build",
            "Migrations: docker compose exec web flask db upgrade",
            "Keycloak: http://localhost:8080",
            "Flask app: http://localhost:5000",
        ],
    )
    return doc


def build_erick_doc() -> Document:
    doc = Document()
    setup_styles(doc)
    add_title(
        doc,
        "manage-system — Half B: Quality, Observability & Delivery",
        "Presenter: Erick De la Rosa | Weeks 5–6 + Day 13 | Prove, observe, and ship the application",
    )

    add_h(doc, "1. Executive Summary")
    add_p(
        doc,
        "This document covers validation and operations: testing pyramid, security scanning, k6 performance, "
        "LGTM observability stack, OpenTelemetry, Grafana dashboards, alerting, and CI/CD.",
    )

    add_h(doc, "2. TODO Days Covered (Erick)")
    add_bullets(
        doc,
        [
            "Day 13: Unit tests (pytest)",
            "Day 29: Schemathesis contract tests",
            "Days 30–31: Playwright E2E",
            "Day 32: OWASP ZAP, pip-audit",
            "Day 33: k6 load/stress/smoke",
            "Day 34: Data/migration tests, exploratory charters",
            "Days 36–38: Observability, dashboards, alerts",
            "Days 39–40: GitHub Actions, Jenkins, verify scripts",
        ],
    )

    add_h(doc, "3. Testing Pyramid")
    add_bullets(
        doc,
        [
            "Unit — business logic (tests/test_products.py, test_stock.py)",
            "Contract — OpenAPI compliance (tests/test_contract.py)",
            "Data — Postgres migrations and constraints (tests/data/)",
            "E2E — Playwright browser flows (tests/e2e/)",
            "Performance — k6 (tests/k6/)",
            "Security — ZAP + pip-audit",
        ],
    )

    add_page_break(doc)
    add_h(doc, "4. Unit Tests (Day 13)")
    add_p(
        doc,
        "Files: tests/test_products.py (14 tests), tests/test_stock.py (13 tests). "
        "tests/conftest.py sets DATABASE_URL=sqlite:///:memory: and TESTING=True to bypass JWT.",
    )
    add_bullets(
        doc,
        [
            "Product CRUD: create, update, delete, duplicate SKU rejection",
            "Product search: filter by name, category, pagination, sort",
            "Stock: entry/exit movements, insufficient stock rejection",
            "Stock alerts: products at or below min_stock flagged",
        ],
    )
    add_p(doc, "Run: pytest tests/test_products.py tests/test_stock.py -v")

    add_h(doc, "5. Contract Tests — Schemathesis (Day 29)")
    add_code(
        doc,
        "tests/test_contract.py:",
        read_snippet("tests/test_contract.py", 14, 39)
        + "\n\n"
        + read_snippet("tests/test_contract.py", 69, 73),
    )

    add_h(doc, "6. Data and Migration Tests (Day 34)")
    add_code(
        doc,
        "tests/data/test_migrations.py:",
        read_snippet("tests/data/test_migrations.py", 16, 36),
    )

    add_h(doc, "7. E2E Tests — Playwright (Days 30–31)")
    add_code(
        doc,
        "tests/e2e/keycloak_helpers.py:",
        read_snippet("tests/e2e/keycloak_helpers.py", 1, 36),
    )
    add_p(doc, "Env: E2E_BASE_URL, E2E_ALICE_USER, E2E_MANAGER_USER, E2E_*_PASSWORD")

    add_p(doc, "Four E2E tests in tests/e2e/: test_auth, test_products, test_stock, test_audit.")
    add_p(doc, "docker-compose.ci.yml runs flask db upgrade on startup; prepare-keycloak-e2e.sh fixes kratos_boss profile.")

    add_page_break(doc)
    add_h(doc, "8. Security Testing (Day 32)")
    add_bullets(
        doc,
        [
            "docker compose --profile security run zap",
            "pip-audit -r requirements.txt",
            "docs/security/day32-findings.md",
        ],
    )

    add_h(doc, "9. Performance — k6 (Day 33)")
    add_code(
        doc,
        "tests/k6/auth.js — Keycloak password grant:",
        read_snippet("tests/k6/auth.js", 1, 37),
    )
    add_bullets(
        doc,
        [
            "tests/k6/auth.js — Keycloak token grant",
            "tests/k6/load-test.js — 50 VUs GET /api/products",
            "tests/k6/stress-test.js — POST /api/stock/movement",
            "tests/k6/smoke-test.js — CI quick check",
        ],
    )

    add_h(doc, "10. Observability (Days 36–38)")
    add_code(
        doc,
        "app/telemetry.py — metrics and OTEL:",
        read_snippet("app/telemetry.py", 30, 67),
    )
    add_bullets(
        doc,
        [
            "Aplicación — latency, throughput, errors",
            "Infraestructura — CPU, memory, DB pool",
            "Negocio — products_created_total, stock_movements_total",
            "Seguridad — auth_failures_total, invalid_tokens_total",
        ],
    )
    add_code(
        doc,
        "observability/prometheus/rules/alerts.rules.yml:",
        read_snippet("observability/prometheus/rules/alerts.rules.yml", 1, 14),
    )

    add_p(doc, "OTEL env: OTEL_ENABLED=true, OTEL_EXPORTER_OTLP_ENDPOINT=http://alloy:4318/v1/traces")
    add_p(doc, "Prometheus scrapes web:5000/metrics and node-exporter:9100 per observability/prometheus/prometheus.yml")

    add_h(doc, "10b. Exploratory Testing (Day 34)")
    add_p(doc, "docs/exploratory-charters.md defines three 30-minute charters:")
    add_bullets(
        doc,
        [
            "Charter 1 — SKU edge cases (duplicates, unicode, length limits)",
            "Charter 2 — RBAC: can alice_worker access manager routes?",
            "Charter 3 — Data integrity: cascade delete, negative stock, audit log completeness",
        ],
    )

    add_page_break(doc)
    add_h(doc, "11. CI/CD (Days 39–40)")
    add_bullets(
        doc,
        [
            ".github/workflows/ci.yml — build, unit, API, data, coverage, E2E, security, docker",
            "Jenkinsfile — full pipeline + verify-stack + k6 smoke",
            "scripts/run-full-test-suite.sh",
            "scripts/verify-stack.sh",
        ],
    )

    add_h(doc, "12. Demo Script (~12 min)")
    add_bullets(
        doc,
        [
            "pytest unit tests",
            "pytest contract tests",
            "Grafana — 4 dashboards",
            "bash scripts/verify-stack.sh",
            "Green CI pipeline + k6 smoke",
        ],
    )

    add_h(doc, "13. Bridge to Uziel's Half")
    add_p(
        doc,
        "All tests target Uziel's API, auth middleware, and UI templates. Metrics feed Grafana dashboards. "
        "Build (Uziel) → Verify & Operate (Erick).",
    )

    add_h(doc, "Appendix: Commands")
    add_bullets(
        doc,
        [
            "Full suite: bash scripts/run-full-test-suite.sh",
            "E2E: docker compose -f docker-compose.ci.yml up -d && pytest tests/e2e -m e2e -v",
            "Grafana: http://localhost:3000",
        ],
    )
    return doc


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    uziel_path = OUT / "Uziel-App-Core-Presentation.docx"
    erick_path = OUT / "Erick-Quality-Observability-Presentation.docx"
    build_uziel_doc().save(str(uziel_path))
    build_erick_doc().save(str(erick_path))
    print(f"Created: {uziel_path}")
    print(f"Created: {erick_path}")


if __name__ == "__main__":
    main()
